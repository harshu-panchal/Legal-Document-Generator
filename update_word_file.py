from docx import Document
import os
import json
import datetime
import docx2pdf


class WordDocument:
    
    data_list_for_single_aadhar_num = ["aadhar_number",
                                        "Name_",
                                        "father_name",
                                        "mother_name",
                                        "phone_number",
                                        "email",
                                        "date_of_birth",
                                        "permanent_address",
                                        "local_address",
                                        "age",
                                        "occupation",
                                        "city",
                                        "pincode",
                                        "marriage_status"]
    
    data_list_for_double_aadhar_num = ["aadhar_number1",
                                        "Name_1",
                                        "father_name1",
                                        "mother_name1",
                                        "phone_number1",
                                        "email1",
                                        "date_of_birth1",
                                        "permanent_address1",
                                        "local_address1",
                                        "age1",
                                        "occupation1",
                                        "city1",
                                        "pincode1",
                                        "marriage_status1",
                                        
                                        "aadhar_number2",
                                        "Name_2",
                                        "father_name2",
                                        "mother_name2",
                                        "phone_number2",
                                        "email2",
                                        "date_of_birth2",
                                        "permanent_address2",
                                        "local_address2",
                                        "age2",
                                        "occupation2",
                                        "city2",
                                        "pincode2",
                                        "marriage_status2"]

    input_lst = []
    ch = ''
    ch_for_pdf = ''
    arguments = []
    _date = "_date_"
    file_names = []
    number_of_members = 0



    def get_file_names(self, folder_path = "Raw documents"):
        self.file_names=[]
        for file_name in os.listdir(folder_path):
            if os.path.isfile(os.path.join(folder_path, file_name)):
                self.file_names.append(file_name) 
        return self.file_names


    def choice(self, choice):
        files = self.get_file_names()
        self.ch_for_pdf = choice
        self.ch = choice+".docx"
        print(self.ch)
        if self.ch in files:
            doc_info = json.loads(open("static\\json files\\document_info.json").read())
            doc = doc_info
            self.number_of_members = doc.get(self.ch)
            print("number of members are : ",self.number_of_members)
            if self.number_of_members == 1:
                self.arguments = self.data_list_for_single_aadhar_num
            elif self.number_of_members == 2:
                self.arguments = self.data_list_for_double_aadhar_num
        return self.number_of_members
           

    def take_input_for_single_person(self,aadhar):
        self.input_lst = []
        val=[]
        api_data = json.loads(open("static\\json files\\API_data.json").read())

        for data in range(len(api_data)):
            if api_data[data]["aadhar_number"]==str(aadhar):
                for i in range(len(self.arguments)):
                    val.append(list(str(api_data[data][str(self.arguments[i])])))
       

        for char_list in val:
            self.input_lst.append(''.join(char_list))
        print(self.input_lst)


    
    def take_input_for_two_person(self,aadhar1,aadhar2):
        self.input_lst = []
        val=[]
        api_data = json.loads(open("static\\json files\\API_data.json").read())

        for data in range(len(api_data)):
            if api_data[data]["aadhar_number"]==str(aadhar1):
                for i in range(int(len(self.arguments)/2)):
                    val.append(list(str(api_data[data][str(self.arguments[i][:-1])])))
            if api_data[data]["aadhar_number"]==str(aadhar2):
                for j in range(int(len(self.arguments)/2),len(self.arguments)):
                    val.append(list(str(api_data[data][str(self.arguments[j][:-1])])))

        for char_list in val:
            self.input_lst.append(''.join(char_list))
        print(self.input_lst)

        

    def modify_word_document(self):
        file_path = 'Raw documents\\'+self.ch
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                for i in range(len(self.arguments)):
                    old_text = self.arguments[i]
                    new_text = self.input_lst[i]
                    if self._date in text:
                        today = datetime.datetime.now()
                        formatted_date = today.strftime("%d/%m/%Y")
                        text = text.replace(str(self._date), str(formatted_date))
                    if old_text in text:
                        text = text.replace(str(old_text), str(new_text))
                run.text = text

        path = f"updated documents\\{self.ch}"
        pdf_path = f"static/updated pdf/{self.ch_for_pdf+'.pdf'}"    
        doc.save(path)
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
            docx2pdf.convert(path,pdf_path)   
        else:
            docx2pdf.convert(path,pdf_path)
        return self.ch   



if __name__ == "__main__":
    
    # a = WordDocument()
    # print(a.get_file_names())
    # print(a.get_file_names())
    # a.choice("Bus Undertaking")
    # a.wordToPdf()
    # a.take_input(12353) 
    # a.modify_word_document()
    print("Done")
